import io

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook
from pypdf import PdfWriter

from backend.app.services.documents.acquisition import (
    DEFAULT_MAX_SIZE_BYTES,
    acquire_document,
)
from backend.app.services.documents.errors import (
    DocumentParseError,
    DocumentTooLarge,
    InvalidDocument,
)
from backend.app.services.documents.models import (
    DocumentSource,
    DocumentType,
)


def test_acquire_valid_csv():
    content = b"name,age\nAlice,30\nBob,40\n"

    result = acquire_document(
        filename="data.csv",
        content=content,
    )

    assert result.metadata.filename == "data.csv"
    assert result.metadata.document_type == DocumentType.CSV
    assert result.metadata.media_type == "text/csv"
    assert result.metadata.source == DocumentSource.UPLOAD
    assert result.metadata.size_bytes == len(content)
    assert result.records == [
        {"name": "Alice", "age": "30"},
        {"name": "Bob", "age": "40"},
    ]
    assert result.text is None


def test_acquire_csv_with_utf8_bom():
    result = acquire_document(
        filename="data.csv",
        content=b"\xef\xbb\xbfname,value\nAlice,100\n",
    )

    assert result.records == [
        {"name": "Alice", "value": "100"},
    ]


def test_acquire_csv_missing_header_raises():
    with pytest.raises(InvalidDocument):
        acquire_document(
            filename="data.csv",
            content=b"",
        )


def test_acquire_csv_empty_column_raises():
    with pytest.raises(InvalidDocument):
        acquire_document(
            filename="data.csv",
            content=b"name,,age\nAlice,30,25\n",
        )


def test_acquire_csv_too_many_values_raises():
    with pytest.raises(InvalidDocument):
        acquire_document(
            filename="data.csv",
            content=b"name,age\nAlice,30,extra\n",
        )


def test_acquire_csv_invalid_utf8_raises():
    with pytest.raises(DocumentParseError):
        acquire_document(
            filename="data.csv",
            content=b"name\n\xff\n",
        )


def test_acquire_json_object():
    result = acquire_document(
        filename="data.json",
        content=b'{"name": "Alice", "age": 30}',
    )

    assert result.metadata.document_type == DocumentType.JSON
    assert result.records == [
        {"name": "Alice", "age": 30},
    ]
    assert result.text is None


def test_acquire_json_array():
    result = acquire_document(
        filename="data.json",
        content=b'[{"name": "Alice"}, {"name": "Bob"}]',
    )

    assert result.records == [
        {"name": "Alice"},
        {"name": "Bob"},
    ]


def test_acquire_invalid_json_raises():
    with pytest.raises(DocumentParseError):
        acquire_document(
            filename="data.json",
            content=b'{"name": "Alice"',
        )


def test_acquire_json_non_object_array_raises():
    with pytest.raises(InvalidDocument):
        acquire_document(
            filename="data.json",
            content=b'[{"name": "Alice"}, "invalid"]',
        )


@pytest.mark.parametrize(
    "content",
    [
        b"null",
        b"true",
        b"123",
        b'"text"',
    ],
)
def test_acquire_json_primitive_raises(content):
    with pytest.raises(InvalidDocument):
        acquire_document(
            filename="data.json",
            content=content,
        )


def test_acquire_txt():
    result = acquire_document(
        filename="notes.txt",
        content=b"Hello UWDE\nSecond line",
    )

    assert result.metadata.document_type == DocumentType.TXT
    assert result.metadata.media_type == "text/plain"
    assert result.text == "Hello UWDE\nSecond line"
    assert result.records == []


def test_acquire_txt_with_bom():
    result = acquire_document(
        filename="notes.txt",
        content=b"\xef\xbb\xbfHello UWDE",
    )

    assert result.text == "Hello UWDE"


def test_acquire_txt_invalid_utf8_raises():
    with pytest.raises(DocumentParseError):
        acquire_document(
            filename="notes.txt",
            content=b"\xff",
        )


def test_acquire_html():
    content = b"""
    <html>
      <head><title>UWDE Test</title></head>
      <body>
        <h1>Hello UWDE</h1>
        <p>Document acquisition works.</p>
        <script>ignored()</script>
        <a href="https://example.com">Example</a>
      </body>
    </html>
    """

    result = acquire_document(
        filename="page.html",
        content=content,
    )

    assert result.metadata.document_type == DocumentType.HTML
    assert result.metadata.media_type == "text/html"
    assert "Hello UWDE" in result.text
    assert "Document acquisition works." in result.text
    assert "ignored" not in result.text
    assert result.raw_metadata["title"] == "UWDE Test"
    assert result.raw_metadata["links"] == ["https://example.com"]


def test_acquire_xlsx():
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Data"
    worksheet.append(["name", "score"])
    worksheet.append(["Alice", 95])
    worksheet.append(["Bob", 88])

    buffer = io.BytesIO()
    workbook.save(buffer)

    result = acquire_document(
        filename="data.xlsx",
        content=buffer.getvalue(),
    )

    assert result.metadata.document_type == DocumentType.XLSX
    assert result.metadata.media_type == (
        "application/vnd.openxmlformats-officedocument."
        "spreadsheetml.sheet"
    )
    assert result.records == [
        {"name": "Alice", "score": 95, "_worksheet": "Data"},
        {"name": "Bob", "score": 88, "_worksheet": "Data"},
    ]
    assert result.raw_metadata["worksheets"] == [
        {"name": "Data", "row_count": 2},
    ]


def test_acquire_invalid_xlsx_raises():
    with pytest.raises(DocumentParseError):
        acquire_document(
            filename="data.xlsx",
            content=b"not-an-xlsx-file",
        )


def test_acquire_pdf():
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)

    buffer = io.BytesIO()
    writer.write(buffer)

    result = acquire_document(
        filename="document.pdf",
        content=buffer.getvalue(),
    )

    assert result.metadata.document_type == DocumentType.PDF
    assert result.metadata.media_type == "application/pdf"
    assert result.raw_metadata["page_count"] == 1
    assert result.text == ""


def test_acquire_invalid_pdf_raises():
    with pytest.raises(DocumentParseError):
        acquire_document(
            filename="document.pdf",
            content=b"not-a-pdf",
        )


def test_acquire_docx():
    document = DocxDocument()
    document.add_paragraph("Hello UWDE")
    document.add_paragraph("Document acquisition works.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "95"

    buffer = io.BytesIO()
    document.save(buffer)

    result = acquire_document(
        filename="document.docx",
        content=buffer.getvalue(),
    )

    assert result.metadata.document_type == DocumentType.DOCX
    assert result.metadata.media_type == (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )
    assert "Hello UWDE" in result.text
    assert "Document acquisition works." in result.text
    assert "Name | Score" in result.text
    assert "Alice | 95" in result.text
    assert result.raw_metadata["paragraph_count"] == 2
    assert result.raw_metadata["table_count"] == 1


def test_acquire_invalid_docx_raises():
    with pytest.raises(DocumentParseError):
        acquire_document(
            filename="document.docx",
            content=b"not-a-docx-file",
        )


def test_empty_content_raises():
    with pytest.raises(InvalidDocument):
        acquire_document(
            filename="data.csv",
            content=b"",
        )


def test_document_too_large_raises():
    with pytest.raises(DocumentTooLarge):
        acquire_document(
            filename="data.csv",
            content=b"name\nAlice\n",
            max_size_bytes=1,
        )


def test_invalid_max_size_type_raises():
    with pytest.raises(TypeError):
        acquire_document(
            filename="data.csv",
            content=b"name\nAlice\n",
            max_size_bytes="100",
        )


def test_invalid_max_size_value_raises():
    with pytest.raises(ValueError):
        acquire_document(
            filename="data.csv",
            content=b"name\nAlice\n",
            max_size_bytes=0,
        )


def test_media_type_parameters_are_normalized():
    result = acquire_document(
        filename="data.csv",
        content=b"name\nAlice\n",
        media_type="text/csv; charset=utf-8",
    )

    assert result.metadata.media_type == "text/csv"


def test_extension_takes_priority_over_media_type():
    result = acquire_document(
        filename="data.csv",
        content=b"name\nAlice\n",
        media_type="application/json",
    )

    assert result.metadata.document_type == DocumentType.CSV
    assert result.records == [{"name": "Alice"}]


def test_default_max_size_is_50_mb():
    assert DEFAULT_MAX_SIZE_BYTES == 50 * 1024 * 1024
