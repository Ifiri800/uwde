from __future__ import annotations

import csv
import io
import json
from typing import Any

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from .detection import detect_document_type
from .errors import (
    DocumentParseError,
    DocumentTooLarge,
    InvalidDocument,
)
from .models import (
    DocumentAcquisitionResult,
    DocumentMetadata,
    DocumentSource,
    DocumentType,
)


DEFAULT_MAX_SIZE_BYTES = 50 * 1024 * 1024


_MEDIA_TYPES: dict[DocumentType, str] = {
    DocumentType.CSV: "text/csv",
    DocumentType.XLSX: (
        "application/vnd.openxmlformats-officedocument."
        "spreadsheetml.sheet"
    ),
    DocumentType.JSON: "application/json",
    DocumentType.TXT: "text/plain",
    DocumentType.HTML: "text/html",
    DocumentType.PDF: "application/pdf",
    DocumentType.DOCX: (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ),
}


def _validate_content(
    content: bytes,
    max_size_bytes: int,
) -> None:
    if not isinstance(content, bytes):
        raise TypeError("content must be bytes")

    if not content:
        raise InvalidDocument("Document content cannot be empty.")

    if len(content) > max_size_bytes:
        raise DocumentTooLarge(
            f"Document exceeds maximum size of "
            f"{max_size_bytes} bytes."
        )


def _acquire_csv(
    content: bytes,
) -> list[dict[str, Any]]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError(
            "CSV document is not valid UTF-8."
        ) from exc

    try:
        reader = csv.reader(io.StringIO(text))

        try:
            raw_headers = next(reader)
        except StopIteration:
            raise InvalidDocument(
                "CSV document does not contain a header row."
            )

        fieldnames = [
            field.strip() if isinstance(field, str) else field
            for field in raw_headers
        ]

        if not fieldnames:
            raise InvalidDocument(
                "CSV document does not contain a header row."
            )

        if any(
            not isinstance(field, str) or not field
            for field in fieldnames
        ):
            raise InvalidDocument(
                "CSV document contains an empty column name."
            )

        records: list[dict[str, Any]] = []

        for row in reader:
            if len(row) > len(fieldnames):
                raise InvalidDocument(
                    "CSV document contains more values than columns."
                )

            padded_row = row + [
                None
            ] * (len(fieldnames) - len(row))

            records.append(
                dict(zip(fieldnames, padded_row))
            )

        return records

    except InvalidDocument:
        raise
    except csv.Error as exc:
        raise DocumentParseError(
            f"Failed to parse CSV document: {exc}"
        ) from exc



def _acquire_json(
    content: bytes,
) -> list[dict[str, Any]]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError(
            "JSON document is not valid UTF-8."
        ) from exc

    try:
        payload = json.loads(text)
    except json.JSONDecodeError as exc:
        raise DocumentParseError(
            f"Failed to parse JSON document: {exc.msg}"
        ) from exc

    if isinstance(payload, dict):
        return [payload]

    if isinstance(payload, list):
        if not all(isinstance(item, dict) for item in payload):
            raise InvalidDocument(
                "JSON arrays must contain only objects."
            )

        return payload

    raise InvalidDocument(
        "JSON document must contain an object or an array of objects."
    )


def _acquire_txt(
    content: bytes,
) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError(
            "Text document is not valid UTF-8."
        ) from exc


def _acquire_html(
    content: bytes,
) -> tuple[str, dict[str, Any]]:
    try:
        soup = BeautifulSoup(content, "html.parser")
    except Exception as exc:
        raise DocumentParseError(
            "Failed to parse HTML document."
        ) from exc

    for element in soup(["script", "style", "noscript", "template"]):
        element.decompose()

    text = soup.get_text("\n", strip=True)

    title = soup.title.get_text(strip=True) if soup.title else None

    metadata: dict[str, Any] = {
        "title": title,
        "links": [
            link.get("href")
            for link in soup.find_all("a", href=True)
        ],
    }

    return text, metadata


def _acquire_xlsx(
    content: bytes,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    try:
        workbook = load_workbook(
            filename=io.BytesIO(content),
            read_only=True,
            data_only=True,
        )
    except Exception as exc:
        raise DocumentParseError(
            "Failed to parse XLSX document."
        ) from exc

    records: list[dict[str, Any]] = []
    worksheet_metadata: list[dict[str, Any]] = []

    try:
        for worksheet in workbook.worksheets:
            rows = worksheet.iter_rows(values_only=True)

            try:
                headers = next(rows)
            except StopIteration:
                worksheet_metadata.append(
                    {
                        "name": worksheet.title,
                        "row_count": 0,
                    }
                )
                continue

            normalized_headers = [
                str(value).strip() if value is not None else ""
                for value in headers
            ]

            if not normalized_headers or all(
                not header for header in normalized_headers
            ):
                raise InvalidDocument(
                    f"Worksheet {worksheet.title!r} "
                    "does not contain a header row."
                )

            if any(not header for header in normalized_headers):
                raise InvalidDocument(
                    f"Worksheet {worksheet.title!r} "
                    "contains an empty column name."
                )

            worksheet_rows = 0

            for row in rows:
                if len(row) > len(normalized_headers):
                    raise InvalidDocument(
                        f"Worksheet {worksheet.title!r} "
                        "contains more values than columns."
                    )

                padded = list(row) + [
                    None
                ] * (len(normalized_headers) - len(row))

                record = dict(
                    zip(normalized_headers, padded)
                )
                record["_worksheet"] = worksheet.title

                records.append(record)
                worksheet_rows += 1

            worksheet_metadata.append(
                {
                    "name": worksheet.title,
                    "row_count": worksheet_rows,
                }
            )

    except InvalidDocument:
        raise
    except Exception as exc:
        raise DocumentParseError(
            "Failed while reading XLSX worksheet data."
        ) from exc
    finally:
        workbook.close()

    return records, {
        "worksheets": worksheet_metadata,
    }


def _acquire_pdf(
    content: bytes,
) -> tuple[str, dict[str, Any]]:
    try:
        reader = PdfReader(
            io.BytesIO(content),
            strict=False,
        )
    except Exception as exc:
        raise DocumentParseError(
            "Failed to parse PDF document."
        ) from exc

    page_text: list[str] = []

    try:
        for page in reader.pages:
            text = page.extract_text() or ""
            page_text.append(text.strip())
    except Exception as exc:
        raise DocumentParseError(
            "Failed to extract text from PDF document."
        ) from exc

    text = "\n\n".join(
        page for page in page_text if page
    )

    metadata: dict[str, Any] = {
        "page_count": len(reader.pages),
    }

    if reader.metadata:
        metadata["pdf_metadata"] = {
            key.lstrip("/"): str(value)
            for key, value in reader.metadata.items()
            if value is not None
        }

    return text, metadata


def _acquire_docx(
    content: bytes,
) -> tuple[str, dict[str, Any]]:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError(
            "Failed to parse DOCX document."
        ) from exc

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_rows: list[str] = []

    try:
        for table in document.tables:
            for row in table.rows:
                values = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                if any(values):
                    table_rows.append(
                        " | ".join(values)
                    )
    except Exception as exc:
        raise DocumentParseError(
            "Failed to extract tables from DOCX document."
        ) from exc

    sections: list[str] = []

    if paragraphs:
        sections.append("\n".join(paragraphs))

    if table_rows:
        sections.append("\n".join(table_rows))

    text = "\n\n".join(sections)

    metadata = {
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
    }

    return text, metadata


def acquire_document(
    *,
    filename: str,
    content: bytes,
    media_type: str | None = None,
    max_size_bytes: int = DEFAULT_MAX_SIZE_BYTES,
) -> DocumentAcquisitionResult:
    """
    Safely acquire a supported document.

    Structured formats are normalized into records where practical.
    Text-oriented formats are normalized into text.

    Supported formats:
    CSV, XLSX, JSON, TXT, HTML, PDF, DOCX.
    """

    if not isinstance(max_size_bytes, int):
        raise TypeError("max_size_bytes must be an integer")

    if max_size_bytes <= 0:
        raise ValueError(
            "max_size_bytes must be greater than zero"
        )

    _validate_content(
        content,
        max_size_bytes,
    )

    document_type = detect_document_type(
        filename,
        media_type,
    )

    records: list[dict[str, Any]] = []
    text: str | None = None
    raw_metadata: dict[str, Any] = {}

    if document_type == DocumentType.CSV:
        records = _acquire_csv(content)

    elif document_type == DocumentType.JSON:
        records = _acquire_json(content)

    elif document_type == DocumentType.TXT:
        text = _acquire_txt(content)

    elif document_type == DocumentType.HTML:
        text, raw_metadata = _acquire_html(content)

    elif document_type == DocumentType.XLSX:
        records, raw_metadata = _acquire_xlsx(content)

    elif document_type == DocumentType.PDF:
        text, raw_metadata = _acquire_pdf(content)

    elif document_type == DocumentType.DOCX:
        text, raw_metadata = _acquire_docx(content)

    else:
        raise DocumentParseError(
            f"Acquisition parser for "
            f"{document_type.value!r} "
            "is not implemented."
        )

    resolved_media_type = (
        media_type.split(";", 1)[0].strip().lower()
        if media_type
        else _MEDIA_TYPES[document_type]
    )

    metadata = DocumentMetadata(
        filename=filename.strip(),
        media_type=resolved_media_type,
        document_type=document_type,
        size_bytes=len(content),
        source=DocumentSource.UPLOAD,
    )

    return DocumentAcquisitionResult(
        metadata=metadata,
        records=records,
        text=text,
        raw_metadata=raw_metadata,
    )
